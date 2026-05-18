"""Generate MANUAL.docx — the SpeakSmart User and Technical Manual.

Run from the project root:
    python scripts/generate_manual.py

Requires: python-docx  (install with `uv pip install python-docx`)
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "MANUAL.docx"

MONO_FONT = "Consolas"


def add_paragraph(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def add_code(doc: Document, text: str) -> None:
    """Add a single-line code/command paragraph in monospace."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = MONO_FONT
    run.font.size = Pt(10)


def add_code_block(doc: Document, text: str) -> None:
    """Add a multi-line code block in monospace, preserving line breaks."""
    p = doc.add_paragraph()
    for i, line in enumerate(text.splitlines()):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = MONO_FONT
        run.font.size = Pt(10)


def add_mixed(doc: Document, segments: list[tuple[str, bool]]) -> None:
    """Add a paragraph mixing normal text and monospace.

    Each segment is (text, is_mono).
    """
    p = doc.add_paragraph()
    for text, is_mono in segments:
        run = p.add_run(text)
        if is_mono:
            run.font.name = MONO_FONT
            run.font.size = Pt(10)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_bullet_mixed(doc: Document, segments: list[tuple[str, bool]]) -> None:
    p = doc.add_paragraph(style="List Bullet")
    for text, is_mono in segments:
        run = p.add_run(text)
        if is_mono:
            run.font.name = MONO_FONT
            run.font.size = Pt(10)


def add_table_of_contents(doc: Document) -> None:
    """Insert a Word TOC field. Word will populate it on open / field update."""
    heading = doc.add_paragraph()
    h_run = heading.add_run("Table of Contents")
    h_run.bold = True
    h_run.font.size = Pt(16)

    p = doc.add_paragraph()
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and select “Update Field” to build the table of contents."
    run._r.append(placeholder)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    # Force Word to recalculate fields on open so the TOC populates automatically.
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, mono_cols: set[int] | None = None) -> None:
    mono_cols = mono_cols or set()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for c, val in enumerate(row):
            cells[c].text = ""
            run = cells[c].paragraphs[0].add_run(val)
            if c in mono_cols:
                run.font.name = MONO_FONT
                run.font.size = Pt(10)
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # === Title page ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("SpeakSmart")
    t_run.bold = True
    t_run.font.size = Pt(36)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = subtitle.add_run("User and Technical Manual")
    s_run.font.size = Pt(20)

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tag.add_run(
        "Japanese Pronunciation Analysis & Language Training System"
    )
    tag_run.italic = True
    tag_run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    doc.add_page_break()

    # === Table of Contents ===
    add_table_of_contents(doc)

    # === 1. Introduction ===
    doc.add_heading("1. Introduction", level=1)
    add_paragraph(
        doc,
        "SpeakSmart is a Progressive Web App (PWA) that helps Tourism and Hospitality "
        "Management students at Gordon College (Olongapo City, Philippines) practice Japanese "
        "pronunciation through AI-driven speech analysis. Students record themselves speaking "
        "Japanese phrases and receive objective, real-time accuracy scores along with phoneme-level "
        "feedback, while their instructors monitor class-wide performance and identify learners who "
        "need extra support.",
    )
    add_paragraph(doc, "SpeakSmart addresses four concrete problems:")
    add_bullet(doc, "Manual pronunciation evaluation does not scale and is inconsistent across instructors.")
    add_bullet(doc, "Generic language tools are not tuned to the difficulties Filipino learners face with Japanese.")
    add_bullet(doc, "There is no centralized way to track and visualize student progress over time.")
    add_bullet(doc, "Students lack confidence when speaking with Japanese tourists in real settings.")

    # === 2. System Overview ===
    doc.add_heading("2. System Overview", level=1)
    add_paragraph(
        doc,
        "The platform has two user roles. Each role sees a different navigation and dashboard:",
    )
    add_table(
        doc,
        headers=["Role", "Primary Capabilities"],
        rows=[
            [
                "Student",
                "Practice pronunciation, complete graded assignments, view personal progress, earn certificates.",
            ],
            [
                "Instructor",
                "Create classes, design exercises, monitor class analytics, review and grade student submissions.",
            ],
        ],
    )
    add_paragraph(
        doc,
        "SpeakSmart is a mobile-first PWA — students can install it on a phone home screen and use "
        "it like a native app. A modern desktop or laptop browser is also supported for instructors.",
    )

    # ====================================================================
    # PART I — USER MANUAL
    # ====================================================================
    doc.add_heading("Part I — User Manual", level=1)

    # === 3. Getting Started ===
    doc.add_heading("3. Getting Started", level=1)

    doc.add_heading("3.1 Creating an Account", level=2)
    add_paragraph(
        doc,
        "Open SpeakSmart in your browser. From the landing page, tap Sign Up. You can register with "
        "an email address and password or with a Google account. After registering, you will be "
        "asked to complete your profile.",
    )

    doc.add_heading("3.2 Completing Your Profile", level=2)
    add_paragraph(
        doc,
        "On first login the app opens the Complete Profile screen, where you choose your role "
        "(Student or Instructor) and set your display name. Your role determines which navigation "
        "and features are available; it cannot be changed casually after this step.",
    )

    doc.add_heading("3.3 Joining a Class (Students)", level=2)
    add_paragraph(
        doc,
        "Open the Classes page from the navigation. Enter the join code provided by your instructor "
        "and confirm. Once joined, your attempts and submissions will be linked to that class so your "
        "instructor can see your progress.",
    )

    doc.add_heading("3.4 Creating a Class (Instructors)", level=2)
    add_paragraph(
        doc,
        "Open the Classes page and select Create Class. Give the class a name. SpeakSmart generates "
        "a unique join code that you can share with your students. You can regenerate the code if it "
        "is shared accidentally.",
    )

    # === 4. Student Guide ===
    doc.add_heading("4. Student Guide", level=1)

    doc.add_heading("4.1 Home", level=2)
    add_paragraph(
        doc,
        "The Home dashboard summarizes your week: average accuracy across recent attempts, total "
        "phrases practiced, module progress, and your current practice streak.",
    )

    doc.add_heading("4.2 Lessons", level=2)
    add_paragraph(
        doc,
        "Lessons contains six thematic modules covering common tourism and hospitality situations:",
    )
    add_bullet(doc, "Greetings")
    add_bullet(doc, "Hotel")
    add_bullet(doc, "Directions")
    add_bullet(doc, "Food")
    add_bullet(doc, "Emergency")
    add_bullet(doc, "Tour Guide")
    add_paragraph(
        doc,
        "Each module contains several Japanese phrases. Each phrase shows the Japanese script, the "
        "romaji transliteration, and the English translation. Tap a phrase to open the practice view.",
    )

    doc.add_heading("4.3 Practice", level=2)
    add_paragraph(
        doc,
        "The Practice view is where pronunciation analysis happens. You can play the reference audio, "
        "then tap Record to capture your own attempt. Within a few seconds you will see:",
    )
    add_bullet(doc, "An overall accuracy score from 0 to 100%.")
    add_bullet(doc, "Sub-scores for mora timing, consonant clarity, and vowel pronunciation.")
    add_bullet(doc, "Color-coded phoneme chips highlighting which sounds were strong or weak.")
    add_bullet(doc, "Plain-language feedback suggesting what to improve on the next attempt.")
    add_paragraph(
        doc,
        "All attempts are saved automatically so you and your instructor can review them later.",
    )

    doc.add_heading("4.4 Assignments", level=2)
    add_paragraph(
        doc,
        "Assignments lists the exercises your instructor has assigned to you, along with their due "
        "dates and completion status. Open an assignment to record a submission for each phrase. "
        "Your submission is automatically scored by the system; depending on the exercise, your "
        "instructor may also leave a manual grade and feedback.",
    )

    doc.add_heading("4.5 Results", level=2)
    add_paragraph(
        doc,
        "Results shows your full attempt history with filters by module, score range, and date. "
        "You can listen to any of your past recordings to hear your own progress over time.",
    )

    doc.add_heading("4.6 Progress", level=2)
    add_paragraph(
        doc,
        "Progress visualizes your improvement: a weekly bar chart of average accuracy, a module-by-"
        "module breakdown, and the option to export a PDF progress report.",
    )

    doc.add_heading("4.7 Certificates", level=2)
    add_paragraph(
        doc,
        "When you reach the required mastery level on a module, you unlock a completion certificate "
        "that can be viewed and downloaded.",
    )

    doc.add_heading("4.8 Settings", level=2)
    add_paragraph(
        doc,
        "Settings lets you update your display name, view privacy and terms documents, and sign out "
        "of the app.",
    )

    # === 5. Instructor Guide ===
    doc.add_heading("5. Instructor Guide", level=1)

    doc.add_heading("5.1 Overview", level=2)
    add_paragraph(
        doc,
        "Overview is your dashboard. It shows the class average accuracy, the number of active "
        "students this week, a count of flagged students (those falling behind), and a weekly trend "
        "chart of class accuracy.",
    )

    doc.add_heading("5.2 Students", level=2)
    add_paragraph(
        doc,
        "Students is a sortable, filterable table of everyone in your class. For each student it "
        "shows current accuracy, module progress, attempt counts, and a small weekly trend. On "
        "mobile, filters live in a slide-up sheet; on desktop they appear inline. Tap any student "
        "to open a detailed dialog with their recent attempts and module breakdown.",
    )

    doc.add_heading("5.3 Heatmap", level=2)
    add_paragraph(
        doc,
        "Heatmap is a color-coded grid mapping modules against phoneme categories. It surfaces which "
        "specific sounds (for example, the Japanese ‘r’) the class struggles with the most in each "
        "module, helping you target instruction time.",
    )

    doc.add_heading("5.4 Exercises", level=2)
    add_paragraph(
        doc,
        "Exercises is where you build and assign graded work:",
    )
    add_bullet(doc, "Create an exercise by selecting any combination of phrases from the lesson catalog.")
    add_bullet(doc, "Assign the exercise to a class with an optional due date.")
    add_bullet(doc, "Track completion progress at a glance.")
    add_bullet(doc, "Open submissions to play back the student's audio and review system scores.")
    add_bullet(doc, "Optionally enter a teacher score and written feedback, then release the result to the student.")

    doc.add_heading("5.5 Class Management", level=2)
    add_paragraph(
        doc,
        "From the Classes page you can see your enrolled students, regenerate join codes, and "
        "create additional classes.",
    )

    # === 6. Tips & Troubleshooting ===
    doc.add_heading("6. Tips and Troubleshooting (Users)", level=1)
    add_paragraph(doc, "A few things to check if something is not working as expected:")
    add_bullet(
        doc,
        "Microphone access: SpeakSmart needs permission to use your microphone. If recording does "
        "not start, check your browser's site permissions and re-enable the microphone.",
    )
    add_bullet(
        doc,
        "Audio quality: Record in a quiet room and hold the device about 15-20 cm from your mouth. "
        "Background noise can lower your score even when your pronunciation is correct.",
    )
    add_bullet(
        doc,
        "Sign-in problems: Use the password reset link on the login screen, or sign in with Google "
        "if you used Google originally. Make sure your email is typed correctly.",
    )
    add_bullet(
        doc,
        "Offline use: SpeakSmart is a PWA, so cached pages can load offline; however, pronunciation "
        "scoring requires an internet connection because the audio is analyzed on the server.",
    )

    # ====================================================================
    # PART II — TECHNICAL MANUAL
    # ====================================================================
    doc.add_heading("Part II — Technical Manual", level=1)

    # === 7. Architecture ===
    doc.add_heading("7. Architecture", level=1)
    add_paragraph(
        doc,
        "SpeakSmart is organized as a monorepo with two top-level applications:",
    )
    add_bullet_mixed(doc, [("Frontend/", True), (" — a Vue 3 PWA written in TypeScript.", False)])
    add_bullet_mixed(doc, [("Backend/", True), (" — a FastAPI service written in Python 3.14.", False)])
    add_paragraph(
        doc,
        "Data flow on a typical pronunciation attempt: the browser records audio via the MediaRecorder "
        "API, uploads it through an Axios client to FastAPI, which verifies the spoken phrase via "
        "OpenAI Whisper, scores it using a combination of Azure Pronunciation Assessment and an "
        "in-process DTW + librosa pipeline, persists the result to PostgreSQL, and uploads the audio "
        "to Cloudflare R2 for replay. The frontend uses Pinia stores (auth, attempts, classes, "
        "modules, progress) to cache state, and Vue Router enforces role-based access.",
    )

    # === 8. Tech Stack ===
    doc.add_heading("8. Tech Stack", level=1)

    doc.add_heading("8.1 Frontend", level=2)
    add_bullet(doc, "Vue 3 with the Composition API and <script setup>")
    add_bullet(doc, "TypeScript 5")
    add_bullet(doc, "Vite 8 as the build tool")
    add_bullet(doc, "Pinia 3 for state management")
    add_bullet(doc, "Vue Router 4 for navigation")
    add_bullet(doc, "Tailwind CSS 4 with the Reka UI component library")
    add_bullet(doc, "Axios for HTTP")
    add_bullet(doc, "vite-plugin-pwa for service worker and offline support")
    add_bullet(doc, "Firebase Web SDK for authentication")
    add_bullet(doc, "VueUse, html2canvas, jsPDF, TanStack Vue Table, and Lucide icons")

    doc.add_heading("8.2 Backend", level=2)
    add_bullet(doc, "FastAPI on Python 3.14, served by Uvicorn")
    add_bullet(doc, "SQLAlchemy 2.0 (async) with asyncpg")
    add_bullet(doc, "Alembic for database migrations")
    add_bullet(doc, "Pydantic v2 for settings and request/response schemas")
    add_bullet(doc, "Firebase Admin SDK for ID token verification")
    add_bullet(doc, "slowapi for rate limiting (optional Redis backend)")

    doc.add_heading("8.3 Speech and Audio", level=2)
    add_bullet(doc, "OpenAI Whisper for automatic speech recognition (phrase verification)")
    add_bullet(doc, "Azure Cognitive Services Pronunciation Assessment for phoneme-level scoring")
    add_bullet(doc, "librosa, PyOpenJTalk, Parselmouth, scipy, and a DTW algorithm for mora timing and feature analysis")
    add_bullet(doc, "Google Cloud Text-to-Speech for generating reference audio for phrases")

    doc.add_heading("8.4 Infrastructure", level=2)
    add_bullet(doc, "PostgreSQL (local instance via Docker Compose; managed instance in production)")
    add_bullet(doc, "Cloudflare R2 for audio file storage")
    add_bullet(doc, "Firebase Authentication for identity (Email and Google providers)")

    # === 9. Repository Layout ===
    doc.add_heading("9. Repository Layout", level=1)
    add_paragraph(doc, "Key folders and files:")
    add_code_block(
        doc,
        "SpeakSmart/\n"
        "├── Frontend/\n"
        "│   └── src/\n"
        "│       ├── views/           (student/ and instructor/ pages)\n"
        "│       ├── components/      (shared, instructor, auth, ui, onboarding, landing)\n"
        "│       ├── stores/          (Pinia: auth, attempts, classes, modules, progress)\n"
        "│       ├── api/             (Axios + per-resource API modules)\n"
        "│       ├── router/          (Vue Router with role-based guards)\n"
        "│       ├── composables/     (useAudioRecorder, useAuthRedirect, ...)\n"
        "│       ├── types/\n"
        "│       └── firebase.ts\n"
        "├── Backend/\n"
        "│   ├── app/\n"
        "│   │   ├── api/v1/endpoints/  (auth, users, classes, modules, phrases,\n"
        "│   │   │                       attempts, exercises, progress, analytics)\n"
        "│   │   ├── db/models/         (user, class_, module, phrase, attempt,\n"
        "│   │   │                       exercise, progress, analytics)\n"
        "│   │   ├── services/          (asr, pronunciation, phoneme_assessment,\n"
        "│   │   │                       pitch_accent, scoring, verification,\n"
        "│   │   │                       firebase_auth, r2_storage, ...)\n"
        "│   │   ├── schemas/\n"
        "│   │   ├── core/\n"
        "│   │   ├── config.py\n"
        "│   │   └── main.py\n"
        "│   ├── alembic/               (database migrations)\n"
        "│   └── docker-compose.yaml    (local Postgres + Adminer)\n"
        "└── scripts/                   (utility scripts, including this generator)",
    )

    # === 10. Setup ===
    doc.add_heading("10. Setup and Installation", level=1)

    doc.add_heading("10.1 Prerequisites", level=2)
    add_bullet(doc, "Node.js 20+ and npm")
    add_bullet(doc, "Python 3.14 and uv (https://github.com/astral-sh/uv)")
    add_bullet(doc, "Docker Desktop (optional, for the local PostgreSQL container)")
    add_bullet(doc, "A Firebase project with Email/Password and Google sign-in enabled")
    add_bullet(doc, "A Cloudflare R2 bucket and access keys")
    add_bullet(doc, "An Azure Speech resource (key + region)")
    add_bullet(doc, "A Google Cloud service account JSON for Text-to-Speech")

    doc.add_heading("10.2 Frontend", level=2)
    add_code_block(
        doc,
        "cd Frontend\n"
        "npm install\n"
        "npm run dev        # http://localhost:5173\n"
        "npm run build      # production build (outputs dist/)\n"
        "npm run preview    # preview the production build locally",
    )

    doc.add_heading("10.3 Backend", level=2)
    add_code_block(
        doc,
        "cd Backend\n"
        "uv python install 3.14\n"
        "uv sync                              # install dependencies from pyproject.toml\n"
        "docker compose up -d                 # start local Postgres on :5432\n"
        "uv run alembic upgrade head          # apply database migrations\n"
        "uv run uvicorn app.main:app --reload # http://localhost:8000",
    )

    # === 11. Environment Variables ===
    doc.add_heading("11. Environment Variables", level=1)

    doc.add_heading("11.1 Frontend (Frontend/.env)", level=2)
    add_table(
        doc,
        headers=["Variable", "Purpose"],
        rows=[
            ["VITE_FIREBASE_API_KEY", "Firebase Web API key"],
            ["VITE_FIREBASE_AUTH_DOMAIN", "Firebase auth domain"],
            ["VITE_FIREBASE_PROJECT_ID", "Firebase project ID"],
            ["VITE_FIREBASE_STORAGE_BUCKET", "Firebase storage bucket"],
            ["VITE_FIREBASE_MESSAGING_SENDER_ID", "Firebase messaging sender ID"],
            ["VITE_FIREBASE_APP_ID", "Firebase app ID"],
            ["VITE_MEASUREMENT_ID", "Firebase Analytics measurement ID (optional)"],
            ["VITE_API_BASE_URL", "Base URL of the FastAPI backend (e.g. http://127.0.0.1:8000)"],
        ],
        mono_cols={0},
    )

    doc.add_heading("11.2 Backend (Backend/.env)", level=2)
    add_table(
        doc,
        headers=["Variable", "Purpose"],
        rows=[
            ["APP_ENV", "Environment label (development / production)"],
            ["SECRET_KEY", "Application secret"],
            ["DATABASE_URL", "Async PostgreSQL DSN (postgresql+asyncpg://...)"],
            ["FIREBASE_PROJECT_ID", "Firebase project for token verification"],
            ["FIREBASE_PRIVATE_KEY / CLIENT_EMAIL / ...", "Firebase Admin SDK credentials"],
            ["R2_ACCOUNT_ID, R2_ACCESS_KEY_ID, R2_SECRET_ACCESS_KEY", "Cloudflare R2 credentials"],
            ["R2_BUCKET_NAME, R2_PUBLIC_URL", "R2 bucket and its public-facing base URL"],
            ["GOOGLE_APPLICATION_CREDENTIALS", "Path to the Google Cloud service-account JSON"],
            ["AZURE_SPEECH_KEY, AZURE_SPEECH_REGION", "Azure Pronunciation Assessment credentials"],
            ["MAX_AUDIO_SIZE_MB", "Maximum upload size for an audio attempt (default 10)"],
            ["ASR_PROVIDER", "Speech recognition provider (default openai_whisper)"],
            ["OPENAI_WHISPER_MODEL", "Whisper model name (small, turbo, large, ...)"],
            ["OPENAI_WHISPER_DEVICE", "auto / cpu / cuda"],
            ["OPENAI_WHISPER_LANGUAGE", "Forced ASR language (default ja)"],
            ["BACKEND_CORS_ORIGINS", "Comma-separated allowed origins for CORS"],
            ["REDIS_URL", "Optional Redis DSN for distributed rate limiting"],
            ["RATE_LIMIT_* (multiple)", "slowapi rate-limit strings per endpoint tier"],
        ],
        mono_cols={0},
    )

    # === 12. Authentication & Authorization ===
    doc.add_heading("12. Authentication and Authorization", level=1)
    add_paragraph(
        doc,
        "Authentication is handled by Firebase Authentication. The high-level flow is:",
    )
    add_bullet(doc, "The browser uses the Firebase Web SDK to sign the user in with email/password or Google.")
    add_bullet_mixed(
        doc,
        [
            ("After sign-in the frontend calls ", False),
            ("POST /api/v1/auth/register", True),
            (" with the chosen role and display name on first login.", False),
        ],
    )
    add_bullet(doc, "The backend verifies the Firebase ID token using the Firebase Admin SDK and upserts a user row.")
    add_bullet_mixed(
        doc,
        [
            ("Every subsequent request includes ", False),
            ("Authorization: Bearer <idToken>", True),
            (", attached automatically by the Axios interceptor in ", False),
            ("Frontend/src/api/axios.ts", True),
            (".", False),
        ],
    )
    add_bullet_mixed(
        doc,
        [
            ("Vue Router guards in ", False),
            ("Frontend/src/router/index.ts", True),
            (" redirect unauthenticated users to /login and enforce role-based access for instructor and student routes.", False),
        ],
    )
    add_bullet(doc, "Backend FastAPI dependencies resolve the current user and role for every protected endpoint.")

    # === 13. API Reference ===
    doc.add_heading("13. API Reference (Grouped)", level=1)
    add_paragraph(
        doc,
        "All endpoints are mounted under /api/v1. The groups below mirror the modules in "
        "Backend/app/api/v1/endpoints/.",
    )
    add_table(
        doc,
        headers=["Group", "Representative Endpoints", "Purpose"],
        rows=[
            ["auth", "POST /auth/register, GET /auth/me, GET /auth/verify", "Profile creation and token validation"],
            ["users", "GET /users/me, PATCH /users/me", "Read and update the current user's profile"],
            ["classes", "GET/POST /classes, join by code, regenerate code", "Class CRUD and enrollment"],
            ["modules", "GET /modules", "Read-only module catalog"],
            ["phrases", "GET /phrases", "Read-only phrase catalog with reference audio URLs"],
            ["attempts", "POST /attempts (multipart audio), GET /attempts, GET /attempts/{id}", "Submit and retrieve pronunciation attempts"],
            ["exercises", "GET/POST /exercises, assignment and submission routes", "Instructor-created graded work"],
            ["progress", "GET /progress", "Per-student progress summaries"],
            ["analytics", "GET /analytics/* (overview, students, heatmap, weekly trends)", "Instructor dashboard data"],
        ],
        mono_cols={1},
    )

    # === 14. Database Schema ===
    doc.add_heading("14. Database Schema (Summary)", level=1)
    add_paragraph(
        doc,
        "PostgreSQL tables are modeled in Backend/app/db/models/. Key tables:",
    )
    add_table(
        doc,
        headers=["Table", "Purpose"],
        rows=[
            ["users", "Students and instructors; keyed by Firebase UID, stores role and display name."],
            ["classes", "Instructor-owned classes with unique join codes."],
            ["class_memberships", "Many-to-many student ↔ class enrollment."],
            ["modules", "The six lesson modules and their ordering."],
            ["phrases", "Individual phrases (Japanese, romaji, English) plus reference audio URL."],
            ["attempts", "Every pronunciation attempt: audio URL, accuracy/mora/consonant/vowel scores, phoneme error map, feedback text, verification status."],
            ["progress_summary", "Per-student, per-module pre-aggregated statistics for fast dashboard loads."],
            ["exercises", "Instructor-built exercise definitions."],
            ["exercise_phrases", "Phrases included in each exercise."],
            ["exercise_assignments", "Which students are assigned which exercise."],
            ["exercise_submissions", "Student submissions, system scores, and optional teacher grade/feedback."],
            ["class_analytics", "Weekly pre-computed class statistics for the instructor Overview."],
        ],
        mono_cols={0},
    )

    # === 15. Audio / Speech Pipeline ===
    doc.add_heading("15. Audio and Speech Pipeline", level=1)
    add_paragraph(
        doc,
        "The pronunciation scoring path is the most distinctive part of the system. End to end:",
    )
    add_bullet_mixed(
        doc,
        [
            ("The browser captures audio with the ", False),
            ("useAudioRecorder", True),
            (" composable and uploads it as multipart/form-data to ", False),
            ("POST /api/v1/attempts", True),
            (".", False),
        ],
    )
    add_bullet(
        doc,
        "Verification: OpenAI Whisper transcribes the clip and confirms the student actually spoke "
        "the target phrase. If the transcript does not match, the attempt is flagged as wrong_phrase "
        "rather than scored as a bad pronunciation.",
    )
    add_bullet(
        doc,
        "Phoneme assessment: Azure Pronunciation Assessment returns per-phoneme accuracy scores.",
    )
    add_bullet(
        doc,
        "Mora-timing and acoustic features: librosa, PyOpenJTalk, and a Dynamic Time Warping (DTW) "
        "comparison against the reference audio yield mora-timing, consonant, and vowel sub-scores.",
    )
    add_bullet(
        doc,
        "A composite accuracy score and human-readable feedback are produced and saved to the "
        "attempts table.",
    )
    add_bullet_mixed(
        doc,
        [
            ("The raw audio file is uploaded to Cloudflare R2; its public URL is stored on ", False),
            ("attempts.audio_file_url", True),
            (" for later playback.", False),
        ],
    )
    add_bullet(
        doc,
        "Reference audio for each phrase is generated ahead of time with Google Cloud Text-to-Speech "
        "and served from R2.",
    )

    # === 16. Build & Deployment ===
    doc.add_heading("16. Build and Deployment", level=1)
    add_paragraph(doc, "Frontend:")
    add_bullet_mixed(doc, [("Run ", False), ("npm run build", True), (" to produce a production bundle in ", False), ("Frontend/dist/", True), (".", False)])
    add_bullet_mixed(doc, [("vite-plugin-pwa emits the service worker and ", False), ("manifest.webmanifest", True), (" so the build is installable as a PWA.", False)])
    add_bullet(doc, "Deploy the dist folder to any static host (Cloudflare Pages, Netlify, Vercel, Nginx, etc.).")
    add_paragraph(doc, "Backend:")
    add_bullet_mixed(doc, [("Containerize using ", False), ("Backend/Dockerfile", True), (" and run with Uvicorn (or any ASGI server).", False)])
    add_bullet_mixed(doc, [("Apply database migrations on each deploy: ", False), ("alembic upgrade head", True), (".", False)])
    add_bullet(doc, "Provide all required environment variables via your hosting platform's secret manager.")

    # === 17. Troubleshooting (Devs) ===
    doc.add_heading("17. Troubleshooting (Developers)", level=1)
    add_bullet_mixed(
        doc,
        [
            ("CORS errors: check ", False),
            ("BACKEND_CORS_ORIGINS", True),
            (" in the backend .env and the allowed origins list in ", False),
            ("Backend/app/main.py", True),
            (".", False),
        ],
    )
    add_bullet(
        doc,
        "401 Unauthorized: usually a Firebase token expiry or a mismatch between the project ID in "
        "the frontend and backend environments. Confirm both reference the same Firebase project.",
    )
    add_bullet_mixed(
        doc,
        [
            ("Audio upload fails: confirm the file is within ", False),
            ("MAX_AUDIO_SIZE_MB", True),
            (" and that the R2 credentials and bucket name are correct.", False),
        ],
    )
    add_bullet_mixed(
        doc,
        [
            ("High scoring latency: tune ", False),
            ("OPENAI_WHISPER_MODEL", True),
            (" (smaller = faster), pick a closer ", False),
            ("AZURE_SPEECH_REGION", True),
            (", or enable GPU via ", False),
            ("OPENAI_WHISPER_DEVICE=cuda", True),
            (".", False),
        ],
    )
    add_bullet_mixed(
        doc,
        [
            ("Migration drift: run ", False),
            ("uv run alembic upgrade head", True),
            ("; if state is inconsistent, inspect ", False),
            ("alembic_version", True),
            (" in the database.", False),
        ],
    )

    # === 18. Glossary ===
    doc.add_heading("18. Glossary", level=1)
    add_table(
        doc,
        headers=["Term", "Meaning"],
        rows=[
            ["Mora", "The rhythmic unit of Japanese pronunciation; roughly one beat per kana character."],
            ["Phoneme", "An individual speech sound (e.g. /k/, /a/) used as the unit of pronunciation scoring."],
            ["DTW", "Dynamic Time Warping — an algorithm that aligns two audio signals of different lengths to compare them."],
            ["ASR", "Automatic Speech Recognition — converting spoken audio to text (here, OpenAI Whisper)."],
            ["PWA", "Progressive Web App — a website installable like a native app, with offline capability."],
            ["R2", "Cloudflare R2 — S3-compatible object storage used for audio files."],
            ["FastAPI", "Python web framework used for the backend API."],
            ["Pinia", "Vue's official state management library; replaces Vuex in this project."],
            ["Reka UI", "Headless component library used together with Tailwind CSS for the frontend UI."],
            ["Alembic", "Database migration tool for SQLAlchemy."],
        ],
    )

    # === 19. Contact & Support ===
    doc.add_heading("19. Contact & Support", level=1)
    add_paragraph(
        doc,
        "For questions, technical issues, or feedback about ALESCAN, contact:",
    )
    add_paragraph(doc, "Project Team — Articulate Development Team", bold=True)
    add_mixed(
        doc,
        [
            ("Email: ", False),
            ("speaksmart@gmail.com", True),
        ],
    )

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
